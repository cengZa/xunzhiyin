from __future__ import annotations

import hashlib
import re
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, Twips
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs" / "本科毕设论文模板-论文主体.docx"
OUT = ROOT / "docs" / "generated" / "thesis_word_final_candidate_v11.docx"
CITATION_PATTERN = re.compile(r"\[(?:\d+(?:\s*[-－]\s*\d+)?)(?:\s*,\s*\d+(?:\s*[-－]\s*\d+)?)*\]")
BODY_FONT_SIZE = 12
REFERENCE_FONT_SIZE = 10.5
TABLE_FONT_SIZE = 10.5
CAPTION_FONT_SIZE = 10.5
HEADING_FONT_SIZES = {
    0: 16,
    1: 16,
    2: 15,
    3: 14,
}

SOURCES = [
    ("abstract_cn", ROOT / "docs" / "05_thesis" / "abstract_cn.md"),
    ("abstract_en", ROOT / "docs" / "05_thesis" / "abstract_en.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter1_intro.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter2_related_work.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter3_analysis.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter4_design.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter5_implementation.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter6_test.md"),
    ("chapter", ROOT / "docs" / "05_thesis" / "chapter7_conclusion.md"),
    ("references", ROOT / "docs" / "05_thesis" / "references.md"),
]

TOC_ENTRIES = [
    (1, "中文摘要", "i"),
    (1, "ABSTRACT", "ii"),
    (1, "目    录", "iv"),
    (1, "1 引言", "1"),
    (2, "1.1 项目背景与意义", "1"),
    (2, "1.2 国内外研究现状与同类方案分析", "2"),
    (2, "1.3 本文主要工作内容", "4"),
    (2, "1.4 本文组织结构", "5"),
    (1, "2 校园社交匹配推荐系统相关理论及技术综述", "7"),
    (2, "2.1 推荐系统概述", "7"),
    (2, "2.2 基于内容的推荐", "8"),
    (2, "2.3 TF-IDF 与用户画像", "9"),
    (2, "2.4 协同过滤与本文取舍", "10"),
    (2, "2.5 可解释推荐与大语言模型辅助解释", "10"),
    (2, "2.6 混合推荐与场景化规则重排", "11"),
    (2, "2.7 本文技术路线选择", "11"),
    (2, "2.8 本章小结", "12"),
    (1, "3 校园社交匹配推荐系统需求分析", "13"),
    (2, "3.1 需求分析综述", "13"),
    (2, "3.2 功能性需求", "13"),
    (3, "3.2.1 用户与标签维护需求", "13"),
    (3, "3.2.2 用户画像构建模块需求", "15"),
    (3, "3.2.3 匹配推荐生成模块需求", "16"),
    (3, "3.2.4 推荐解释与用户反馈模块需求", "18"),
    (2, "3.3 非功能性需求", "21"),
    (2, "3.4 本章小结", "22"),
    (1, "4 校园社交匹配推荐系统概要设计", "23"),
    (2, "4.1 系统整体架构设计", "23"),
    (2, "4.2 系统功能模块设计", "24"),
    (2, "4.3 系统核心流程设计", "25"),
    (2, "4.4 数据存储设计", "26"),
    (3, "4.4.1 概念与逻辑模型设计", "26"),
    (3, "4.4.2 物理模型设计", "27"),
    (2, "4.5 Redis 缓存设计", "35"),
    (2, "4.6 本章小结", "36"),
    (1, "5 校园社交匹配推荐系统详细设计与实现", "37"),
    (2, "5.1 用户画像构建模块", "37"),
    (2, "5.2 匹配推荐生成模块", "41"),
    (2, "5.3 推荐解释与用户反馈模块", "44"),
    (2, "5.4 本章小结", "49"),
    (1, "6 校园社交匹配推荐系统测试", "50"),
    (2, "6.1 测试目标", "50"),
    (2, "6.2 测试环境", "50"),
    (2, "6.3 测试用例设计思路", "51"),
    (2, "6.4 功能性需求测试用例", "51"),
    (2, "6.5 非功能性需求测试用例", "54"),
    (2, "6.6 离线评估设计", "55"),
    (2, "6.7 测试结果分析", "58"),
    (2, "6.8 本章小结", "59"),
    (1, "7 总结与展望", "60"),
    (2, "7.1 全文总结", "60"),
    (2, "7.2 系统展望", "61"),
    (1, "参考文献", "63"),
    (1, "致谢", "66"),
]

FIGURE_IMAGES = {
    "图 3-1": ("figures/ch3-1-user-tag-usecase.png", "图 3-1 用户与标签维护用例图"),
    "图 3-2": ("figures/ch3-2-recommendation-usecase.png", "图 3-2 画像构建与推荐生成用例图"),
    "图 3-3": ("figures/ch3-3-explanation-usecase.png", "图 3-3 推荐解释用例图"),
    "图 3-4": ("figures/ch3-4-feedback-usecase.png", "图 3-4 用户反馈与画像更新用例图"),
    "图 4-1": ("figures/ch4-1-system-architecture.png", "图 4-1 系统整体架构图"),
    "图 4-2": ("figures/ch4-2-function-structure.png", "图 4-2 系统功能模块结构图"),
    "图 4-3": ("figures/ch4-3-core-flow.png", "图 4-3 系统核心流程图"),
    "图 4-4": ("figures/ch4-4-conceptual-model.png", "图 4-4 数据存储 ER 图（Peter Chen 表示法）"),
    "图 4-5": ("figures/ch4-5-physical-model.png", "图 4-5 数据存储物理模型示意图"),
    "图 5-1": ("figures/ch5-1-profile-business-flow.png", "图 5-1 用户画像构建模块业务流程图"),
    "图 5-2": ("figures/ch5-2-profile-class-diagram.png", "图 5-2 用户画像构建模块类图"),
    "图 5-3": ("figures/ch5-3-profile-sequence.png", "图 5-3 用户画像生成时序图"),
    "图 5-4": ("figures/ch5-4-home-screenshot.png", "图 5-4 系统首页画像与推荐展示界面"),
    "图 5-5": ("figures/ch5-5-recommendation-business-flow.png", "图 5-5 匹配推荐生成模块业务流程图"),
    "图 5-6": ("figures/ch5-6-recommendation-class-diagram.png", "图 5-6 匹配推荐生成模块类图"),
    "图 5-7": ("figures/ch5-7-recommendation-sequence.png", "图 5-7 匹配推荐生成时序图"),
    "图 5-8": ("figures/ch5-8-pipeline-screenshot.png", "图 5-8 透明链路页面推荐生成阶段界面"),
    "图 5-9": ("figures/ch5-9-explanation-feedback-business-flow.png", "图 5-9 推荐解释与用户反馈模块业务流程图"),
    "图 5-10": ("figures/ch5-10-explanation-feedback-class-diagram.png", "图 5-10 推荐解释与用户反馈模块类图"),
    "图 5-11": ("figures/ch5-11-feedback-sequence.png", "图 5-11 用户反馈更新时序图"),
    "图 5-12": ("figures/ch5-12-feedback-screenshot.png", "图 5-12 推荐解释与反馈展示界面"),
}

FORMULAS = {
    "公式（5-1）": {
        "tokens": ["w(u,t)", " = ", "tf(u,t)", " × ", "idf(t)", " × ", "decay(t)", " × ", "seed(u,t)"],
        "number": "（式5-1）",
    },
    "公式（5-2）": {
        "tokens": ["decay(t)", " = ", "exp(", "-λ", " × ", "days(t)", ")"],
        "number": "（式5-2）",
    },
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, latin: bool = False) -> None:
    run.font.name = "Times New Roman" if latin else "宋体"
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), "宋体")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_heading_font(run, size: float) -> None:
    run.font.name = "黑体"
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:eastAsia"), "黑体")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = False


def add_runs_with_body_citations(paragraph, text: str, bold: bool = False) -> None:
    pos = 0
    for match in CITATION_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, BODY_FONT_SIZE, bold=bold)
        citation = paragraph.add_run(match.group(0))
        set_run_font(citation, BODY_FONT_SIZE, bold=bold)
        citation.font.superscript = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, BODY_FONT_SIZE, bold=bold)


def should_superscript_body_citations(text: str) -> bool:
    return bool(CITATION_PATTERN.search(text)) and not is_non_body_line(text)


def insert_paragraph_before(marker: Paragraph, style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    marker._p.addprevious(new_p)
    paragraph = Paragraph(new_p, marker._parent)
    if style:
        paragraph.style = style
    return paragraph


def insert_table_before(doc: Document, marker: Paragraph, rows: int, cols: int) -> Table:
    table = doc.add_table(rows=rows, cols=cols)
    marker._p.addprevious(table._tbl)
    return table


def clear_between(start: Paragraph, end: Paragraph) -> None:
    cursor = start._p.getnext()
    while cursor is not None and cursor is not end._p:
        next_cursor = cursor.getnext()
        cursor.getparent().remove(cursor)
        cursor = next_cursor


def clear_after(start: Paragraph, doc: Document) -> None:
    sect_pr = doc._body._element.sectPr
    cursor = start._p.getnext()
    while cursor is not None and cursor is not sect_pr:
        next_cursor = cursor.getnext()
        cursor.getparent().remove(cursor)
        cursor = next_cursor


def set_paragraph_format(paragraph, first_line: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.widow_control = True
    if first_line:
        fmt.first_line_indent = Pt(24)


def clear_document_from(doc: Document, first_body_index: int) -> None:
    body = doc._body._element
    children = list(body)
    sect_pr = body.sectPr
    for child in children[first_body_index:]:
        if child is not sect_pr:
            body.remove(child)


def add_page_break(doc: Document, before: Paragraph | None = None) -> None:
    p = insert_paragraph_before(before) if before else doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def normalize_heading(text: str) -> tuple[int, str] | None:
    if text.startswith("# "):
        title = text[2:].strip()
        m = re.match(r"第\s*(\d+)\s*章\s*(.+)", title)
        if m:
            return 1, m.group(2).strip()
        return 0, title
    if text.startswith("## "):
        return 2, strip_heading_number(text[3:].strip())
    if text.startswith("### "):
        return 3, strip_heading_number(text[4:].strip())
    return None


def strip_heading_number(text: str) -> str:
    return re.sub(r"^\d+(?:\.\d+)*\s+", "", text).strip()


def add_heading(doc: Document, level: int, text: str, before: Paragraph | None = None) -> None:
    if level == 0:
        p = insert_paragraph_before(before, "标题名（不入目录）") if before else doc.add_paragraph(style="标题名（不入目录）")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        p = insert_paragraph_before(before, "1级标题") if before else doc.add_paragraph(style="1级标题")
    elif level == 2:
        p = insert_paragraph_before(before, "2级标题") if before else doc.add_paragraph(style="2级标题")
    else:
        p = insert_paragraph_before(before, "3级标题") if before else doc.add_paragraph(style="3级标题")
    p.text = ""
    run = p.add_run(text)
    fmt = p.paragraph_format
    fmt.keep_with_next = True
    fmt.keep_together = True
    fmt.widow_control = True
    set_heading_font(run, HEADING_FONT_SIZES.get(level, 14))


def is_table_start(lines: list[str], idx: int) -> bool:
    if idx + 1 >= len(lines):
        return False
    return lines[idx].strip().startswith("|") and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[idx + 1])


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip().replace("<br/>", "\n") for cell in line.split("|")]


def display_width(text: str) -> int:
    width = 0
    for char in text:
        width += 1 if ord(char) < 128 else 2
    return width


def table_width_twips(doc: Document) -> int:
    section = doc.sections[-1]
    width = section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    return width if width > 0 else 8500


def column_widths(headers: list[str], rows: list[list[str]], total_width: int) -> list[int]:
    columns = len(headers)
    values = [headers] + rows
    weights: list[int] = []
    for col in range(columns):
        max_width = max((display_width(row[col]) for row in values if col < len(row)), default=8)
        weights.append(max(6, min(max_width, 36)))

    min_width = 1500 if columns <= 3 else 1050 if columns == 4 else 760
    min_width = min(min_width, max(1, total_width // columns))
    remaining = max(0, total_width - min_width * columns)
    weight_sum = sum(weights) or columns
    widths = [min_width + int(remaining * weight / weight_sum) for weight in weights]
    widths[-1] += total_width - sum(widths)
    return widths


def set_table_grid(table, widths: list[int], total_width: int) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_width))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    tbl.insert(1, grid)


def set_cell_width(cell, width: int) -> None:
    cell.width = Twips(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width))


def style_table(table, doc: Document, widths: list[int]) -> None:
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_grid(table, widths, table_width_twips(doc))
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                set_paragraph_format(p, first_line=False)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or len(row.cells) <= 3 else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    set_run_font(run, TABLE_FONT_SIZE, bold=(row_idx == 0))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        border = borders.find(qn(tag))
        if border is None:
            border = OxmlElement(tag)
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def set_cell_margins(cell, top: int = 90, bottom: int = 90, left: int = 120, right: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_row_min_height(row, height: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = tr_pr.find(qn("w:trHeight"))
    if tr_height is None:
        tr_height = OxmlElement("w:trHeight")
        tr_pr.append(tr_height)
    tr_height.set(qn("w:val"), str(height))
    tr_height.set(qn("w:hRule"), "atLeast")


def add_after_table_spacing(doc: Document, before: Paragraph | None = None) -> None:
    spacer = insert_paragraph_before(before) if before else doc.add_paragraph()
    fmt = spacer.paragraph_format
    fmt.first_line_indent = None
    fmt.line_spacing = Pt(1)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.widow_control = True


def add_native_markdown_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    before: Paragraph | None = None,
    widths: list[int] | None = None,
    compact: bool = False,
) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    if before:
        before._p.addprevious(table._tbl)
    total_width = table_width_twips(doc)
    widths = widths or column_widths(headers, rows, total_width)
    widths[-1] += total_width - sum(widths)
    try:
        table.style = "Table Grid"
    except Exception:
        pass
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_grid(table, widths, total_width)

    for row_idx, row_values in enumerate([headers] + rows):
        row = table.rows[row_idx]
        set_row_cant_split(row)
        if compact:
            set_row_min_height(row, 320 if row_idx == 0 else 270)
        else:
            set_row_min_height(row, 460 if row_idx == 0 else 430)
        merge_tail = (
            row_idx != 0
            and len(headers) == 4
            and len(row_values) >= 4
            and row_values[1]
            and row_values[1] == row_values[2] == row_values[3]
        )
        if merge_tail:
            row.cells[1].merge(row.cells[3])
            fill_values = [(0, row_values[0], widths[0]), (1, row_values[1], sum(widths[1:4]))]
        else:
            fill_values = [(col_idx, row_values[col_idx] if col_idx < len(row_values) else "", widths[col_idx]) for col_idx in range(len(headers))]
        for col_idx, text, width in fill_values:
            cell = row.cells[col_idx]
            set_cell_width(cell, width)
            set_cell_borders(cell)
            if compact:
                set_cell_margins(cell, top=35, bottom=35, left=90, right=90)
            else:
                set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, "D9E3F0")
            p = cell.paragraphs[0]
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if row_idx != 0 and (merge_tail or col_idx == len(headers) - 1) else WD_ALIGN_PARAGRAPH.CENTER
            fmt = p.paragraph_format
            fmt.first_line_indent = None
            fmt.left_indent = None
            fmt.right_indent = None
            fmt.space_before = Pt(0 if compact else 1)
            fmt.space_after = Pt(0 if compact else 1)
            fmt.line_spacing = Pt(12 if compact else 14)
            fmt.keep_with_next = row_idx < len(rows)
            fmt.keep_together = True
            run = p.add_run(normalize_inline_text(text))
            set_run_font(run, TABLE_FONT_SIZE, bold=(row_idx == 0))
    add_after_table_spacing(doc, before=before)


def add_table_row_text(doc: Document, text: str, bold_label: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Pt(21)
    p.text = ""
    run = p.add_run(text)
    set_run_font(run, TABLE_FONT_SIZE, bold_label)


def add_table_placeholder_line(doc: Document, text: str, before: Paragraph | None = None, bold: bool = False) -> None:
    p = insert_paragraph_before(before) if before else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(18)
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Pt(21)
    p.text = ""
    run = p.add_run(text)
    set_run_font(run, TABLE_FONT_SIZE, bold)


def add_markdown_table(doc: Document, lines: list[str], idx: int, before: Paragraph | None = None) -> int:
    start_idx = idx
    headers = split_table_row(lines[idx])
    rows: list[list[str]] = []
    idx += 2
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        rows.append(split_table_row(lines[idx]))
        idx += 1

    normalized_headers = [normalize_inline_text(h) for h in headers]
    if normalized_headers == ["序号", "字段名", "类型", "属性", "描述"]:
        add_native_markdown_table(
            doc,
            normalized_headers,
            [[normalize_inline_text(cell) for cell in row[: len(headers)]] for row in rows],
            before=before,
            widths=[560, 2100, 1500, 1800, table_width_twips(doc) - 5960],
        )
        return idx

    if normalized_headers == ["环境项", "配置或用途", "说明"]:
        total = table_width_twips(doc)
        widths = [1300, 3000, total - 4300]
        add_native_markdown_table(
            doc,
            normalized_headers,
            [[normalize_inline_text(cell) for cell in row[: len(headers)]] for row in rows],
            before=before,
            widths=widths,
            compact=True,
        )
        return idx

    if len(normalized_headers) == 4 and rows and rows[0][0] == "用例编号":
        widths = [1100, 2850, 1100, table_width_twips(doc) - 5050]
    else:
        widths = column_widths(normalized_headers, rows, table_width_twips(doc))
    add_native_markdown_table(
        doc,
        normalized_headers,
        [[normalize_inline_text(cell) for cell in row[: len(headers)]] for row in rows],
        before=before,
        widths=widths,
    )
    return idx


def add_placeholder(doc: Document, text: str, before: Paragraph | None = None) -> None:
    p = insert_paragraph_before(before) if before else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = p.paragraph_format
    fmt.line_spacing = Pt(20)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    p.text = ""
    run = p.add_run(f"【{text}】")
    set_run_font(run, CAPTION_FONT_SIZE, True)


def add_figure(doc: Document, figure_ref: str, before: Paragraph | None = None) -> bool:
    image_info = FIGURE_IMAGES.get(figure_ref)
    if not image_info:
        return False
    rel_path, caption = image_info
    image_path = ROOT / "docs" / "generated" / rel_path
    if not image_path.exists():
        return False

    p = insert_paragraph_before(before) if before else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))

    add_paragraph_text(doc, caption, before=before)
    return True


def formula_reference(text: str) -> str | None:
    match = re.fullmatch(r"(?:公式|式)\s*（\s*(\d+)\s*[-－]\s*(\d+)\s*）", text.strip())
    if not match:
        return None
    return f"公式（{match.group(1)}-{match.group(2)}）"


def make_omml_formula(tokens: list[str]):
    runs = []
    for token in tokens:
        space = ' xml:space="preserve"' if token.startswith(" ") or token.endswith(" ") else ""
        runs.append(
            f"""
            <m:r>
              <w:rPr>
                <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math" w:eastAsia="宋体"/>
                <w:sz w:val="21"/>
              </w:rPr>
              <m:t{space}>{escape(token)}</m:t>
            </m:r>
            """
        )
    xml = f"""
    <m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
             xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      {''.join(runs)}
    </m:oMath>
    """
    return parse_xml(xml)


def add_formula(doc: Document, formula_ref: str, before: Paragraph | None = None) -> bool:
    formula = FORMULAS.get(formula_ref)
    if not formula:
        return False
    p = insert_paragraph_before(before) if before else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.right_indent = Pt(21)
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    p.paragraph_format.tab_stops.clear_all()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(3.05), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.15), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run().add_tab()
    p._p.append(make_omml_formula(formula["tokens"]))
    p.add_run().add_tab()
    number_run = p.add_run(formula["number"])
    set_run_font(number_run, CAPTION_FONT_SIZE, latin=False)
    return True


def add_paragraph_text(doc: Document, text: str, style: str | None = None, before: Paragraph | None = None) -> None:
    text = normalize_inline_text(text)
    text = re.sub(r"^(\d+)\.\s+", lambda m: f"（{m.group(1)}）", text)
    p = insert_paragraph_before(before, style) if before else (doc.add_paragraph(style=style) if style else doc.add_paragraph())
    if text.startswith("["):
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_centered_caption(text) else WD_ALIGN_PARAGRAPH.JUSTIFY
    set_paragraph_format(p, first_line=not is_non_body_line(text))
    if text.startswith("["):
        p.paragraph_format.left_indent = Pt(22)
        p.paragraph_format.first_line_indent = Pt(-22)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(20)
    if text.startswith("表 "):
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
    elif text.startswith("表名："):
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
    if text.startswith("图 "):
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_together = True
    p.text = ""
    bold = text.startswith("表名：")
    if should_superscript_body_citations(text):
        add_runs_with_body_citations(p, text, bold=bold)
    else:
        run = p.add_run(text)
        latin = bool(text) and sum(1 for c in text if ord(c) < 128) / max(len(text), 1) > 0.8
        if text.startswith("["):
            size = REFERENCE_FONT_SIZE
        elif text.startswith(("表 ", "表名：", "图 ")):
            size = CAPTION_FONT_SIZE
        else:
            size = BODY_FONT_SIZE
        set_run_font(run, size, bold=bold, latin=latin)


def add_static_toc(doc: Document, before: Paragraph | None = None) -> None:
    add_heading(doc, 0, "目    录", before=before)
    for level, title, page in TOC_ENTRIES:
        p = insert_paragraph_before(before) if before else doc.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        fmt = p.paragraph_format
        fmt.left_indent = Pt(0)
        fmt.first_line_indent = Pt(0 if level == 1 else 12 if level == 2 else 24)
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = Pt(20)
        fmt.tab_stops.clear_all()
        fmt.tab_stops.add_tab_stop(Inches(6.15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        run = p.add_run(f"{title}\t{page}")
        if level == 1:
            run.font.name = "黑体"
            r_fonts = run._element.rPr.rFonts
            r_fonts.set(qn("w:eastAsia"), "黑体")
            r_fonts.set(qn("w:ascii"), "Times New Roman")
            r_fonts.set(qn("w:hAnsi"), "Times New Roman")
            run.font.size = Pt(12)
            run.bold = False
        else:
            set_run_font(run, 12, bold=False)


def normalize_inline_text(text: str) -> str:
    return text.replace("`", "")


def is_centered_caption(text: str) -> bool:
    return text.startswith(("表 ", "图 "))


def is_non_body_line(text: str) -> bool:
    return text.startswith(("关键词：", "Keywords:", "[", "表 ", "表名：", "图 "))


def figure_reference(text: str) -> str | None:
    match = re.search(r"图\s*(\d+)\s*[-－]\s*(\d+)", text)
    if not match:
        return None
    return f"图 {match.group(1)}-{match.group(2)}"


def render_markdown(doc: Document, path: Path, doc_kind: str, before: Paragraph | None = None) -> None:
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    in_code = False
    code_buffer: list[str] = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        if line.startswith("```"):
            if in_code:
                add_placeholder(doc, "图表占位：请在 Word 中按模板绘制流程图或示意图", before=before)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        heading = normalize_heading(line)
        if heading:
            level, title = heading
            if doc_kind in {"abstract_cn", "abstract_en", "references"} and level == 0:
                add_heading(doc, 0, title, before=before)
            elif doc_kind == "chapter":
                add_heading(doc, level, title, before=before)
            elif doc_kind == "references":
                if title != "使用说明":
                    add_heading(doc, 0, title, before=before)
            i += 1
            continue
        if is_table_start(lines, i):
            i = add_markdown_table(doc, lines, i, before=before)
            continue
        formula_ref = formula_reference(line)
        if formula_ref:
            if not add_formula(doc, formula_ref, before=before):
                add_paragraph_text(doc, line, before=before)
            i += 1
            continue
        if line.startswith("> 图") or line.startswith("【图"):
            figure_text = line.lstrip("> ").strip("【】")
            ref = figure_reference(figure_text)
            if not ref or not add_figure(doc, ref, before=before):
                add_placeholder(doc, figure_text, before=before)
            i += 1
            continue
        if line.startswith("- "):
            add_paragraph_text(doc, line[2:], before=before)
            i += 1
            continue
        add_paragraph_text(doc, line, before=before)
        i += 1


def update_cover(doc: Document) -> None:
    title_cn = "基于 Spring Boot 的校园社交匹配与可解释推荐系统的设计与实现"
    title_en = "Design and Implementation of Spring Boot-Based Campus Social Matching and Explainable Recommendation System"
    if len(doc.paragraphs) > 7:
        doc.paragraphs[7].text = title_cn
        doc.paragraphs[7].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in doc.paragraphs[7].runs:
            set_run_font(run, 16, True)
    if len(doc.paragraphs) > 9:
        doc.paragraphs[9].text = title_en
        doc.paragraphs[9].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in doc.paragraphs[9].runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    if len(doc.paragraphs) > 84:
        doc.paragraphs[84].text = (
            "本人声明所呈交的毕业论文（设计），题目 基于 Spring Boot 的校园社交匹配与可解释推荐系统的设计与实现 "
            "是本人在指导教师的指导下，独立进行研究工作所取得的成果。尽我所知，除了文中特别加以标注和致谢中所罗列的内容以外，"
            "论文中不包含其他人已经发表或撰写过的研究成果，也不包含为获得北京交通大学或其他教育机构的学位或证书而使用过的材料。"
        )


def set_section_pg_num(section, fmt: str | None = None, start: int | None = None) -> None:
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num)
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)
    elif qn("w:fmt") in pg_num.attrib:
        del pg_num.attrib[qn("w:fmt")]
    if start is not None:
        pg_num.set(qn("w:start"), str(start))
    elif qn("w:start") in pg_num.attrib:
        del pg_num.attrib[qn("w:start")]


def clear_paragraph_content(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def set_footer_static(section, text: str = "") -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, 10.5, latin=bool(re.fullmatch(r"[IVXLCDM]+", text)))


def set_header_static(section, right_text: str = "") -> None:
    section.header.is_linked_to_previous = False
    paragraph = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(5.8), WD_TAB_ALIGNMENT.RIGHT)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    left = paragraph.add_run("北京交通大学毕业设计（论文）")
    set_run_font(left, 10.5)
    paragraph.add_run().add_tab()
    right = paragraph.add_run(right_text)
    set_run_font(right, 10.5)


def set_footer_page_field(section) -> None:
    section.footer.is_linked_to_previous = False
    paragraph = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
    clear_paragraph_content(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_run = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    instr_run.append(instr)
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result_run = OxmlElement("w:r")
    result = OxmlElement("w:t")
    result.text = "1"
    result_run.append(result)
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    paragraph._p.append(instr_run)
    run2 = paragraph.add_run()
    run2._r.append(separate)
    paragraph._p.append(result_run)
    run3 = paragraph.add_run()
    run3._r.append(end)


def normalize_section_page_numbers(doc: Document) -> None:
    # Template section order after clearing: cover/statement sections, Chinese
    # abstract, English abstract, TOC, body, references, acknowledgements.
    for idx, section in enumerate(doc.sections):
        if idx <= 2:
            set_footer_static(section)
    if len(doc.sections) >= 6:
        set_header_static(doc.sections[3], "中文摘要")
        set_header_static(doc.sections[4], "英文摘要")
        set_header_static(doc.sections[5], "目录")
        set_section_pg_num(doc.sections[3], fmt="lowerRoman", start=1)
        set_section_pg_num(doc.sections[4], fmt="lowerRoman")
        set_section_pg_num(doc.sections[5], fmt="lowerRoman")
        set_footer_page_field(doc.sections[3])
        set_footer_page_field(doc.sections[4])
        set_footer_page_field(doc.sections[5])
    if len(doc.sections) >= 7:
        set_header_static(doc.sections[6], "正文")
        set_section_pg_num(doc.sections[6], start=1)
        set_footer_page_field(doc.sections[6])
    if len(doc.sections) >= 8:
        set_header_static(doc.sections[7], "参考文献")
        set_section_pg_num(doc.sections[7])
        set_footer_page_field(doc.sections[7])
    if len(doc.sections) >= 9:
        set_header_static(doc.sections[8], "致谢")
        set_section_pg_num(doc.sections[8])
        set_footer_page_field(doc.sections[8])


def build() -> None:
    doc = Document(TEMPLATE)
    update_cover(doc)

    cn_start = doc.paragraphs[88]
    en_start = doc.paragraphs[111]
    toc_start = doc.paragraphs[134]
    body_start = doc.paragraphs[163]
    ref_start = doc.paragraphs[232]
    ack_start = doc.paragraphs[258]

    clear_between(cn_start, en_start)
    render_markdown(doc, SOURCES[0][1], "abstract_cn", before=en_start)

    clear_between(en_start, toc_start)
    render_markdown(doc, SOURCES[1][1], "abstract_en", before=toc_start)

    clear_between(toc_start, body_start)
    add_static_toc(doc, before=body_start)

    clear_between(body_start, ref_start)
    for idx, (_, source) in enumerate(SOURCES[2:9]):
        if idx > 0:
            add_page_break(doc, before=ref_start)
        render_markdown(doc, source, "chapter", before=ref_start)

    clear_between(ref_start, ack_start)
    render_markdown(doc, SOURCES[9][1], "references", before=ack_start)

    clear_after(ack_start, doc)
    add_heading(doc, 0, "致    谢")

    normalize_section_page_numbers(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
